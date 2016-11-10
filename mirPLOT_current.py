### dependencies gffread from  http://manpages.ubuntu.com/manpages/trusty/man1/gffread.1.html 

import getopt
import sys
import glob
import os
import fileinput
import re
import locale
from docx import Document
from docx.shared import Inches
from docx.shared import Pt 
from docx.shared import RGBColor
from docx.shared import Length
from Bio import SeqIO
###########################################################################################################################################
################## Functions
###########################################################################################################################################
###########################################################################################################################################
def doreportgenomebam(comment):
		print comment
		print("""
			#####################################################################
			#####################################################################
			############						 ############
			############	 start Folding				 ############
			############						 ############
			#####################################################################
			#####################################################################

			"""	)



		os.system("rm -rf temp/structures.txt")

		document = Document()

		document.add_heading('miRNA check', 0)

		style=document.styles["Normal"]
		font=style.font
		font.name="Mono"
		font.size=Pt(8)


		cmd="gffread -g "+Genome+" "+gff3Matures +" -w temp/matures"
		print(cmd)
		os.system(cmd)


		cmd="gffread -g "+Genome+" "+gff3Stars+" -w temp/stars"
		print(cmd)
		os.system(cmd)

		input_file = open("temp/matures")
		maturesdict = SeqIO.to_dict(SeqIO.parse(input_file, "fasta"))

		input_file2 = open("temp/stars")
		starsdict = SeqIO.to_dict(SeqIO.parse(input_file2, "fasta"))



		mirnum=1






		with open(fastaref, 'r') as inF:
			for line in inF:
				if ">" in line:
					print "lineee", line
					precursorname=line.rstrip()

					pictname=output_dir+"plots/"+precursorname.split(">")[1]+'.png'


					if os.path.isfile(pictname) == False:### if picture does ot exist, next candidaate
						print "doesnt exist", pictname 
						flag="next!"
						mirnum=mirnum+2
					else:## if pict exists, continue
						print "exists", pictname
						flag="continue"


						paragraph = document.add_paragraph()
						prename=line.split("_prec")[0]
						name=prename.split(">")[1].rstrip()
						name2=prename.split(">")[1].rstrip()
						name3=name.split("_pre")[0]
						print "3 This is the name file: ", name
						name_file=name+".fa"
						outfile= open(name_file, 'w+')
						outfile.write(line)


						paragraph.add_run(precursorname+"\n").bold = True
			


				if ">" not in line and flag=="continue":
					#print maturesdict
					print "Name YY;",name3


					try:		
						maturename=maturesdict[name3+"_3p"].id
						print "Try;",maturename
						m1=str(maturesdict[name3+"_3p"].seq)
						print "seq;",m1
						starname=starsdict[name3+"_5p*"].id
						s1=str(starsdict[name3+"_5p*"].seq)
						print "star",s1
						starseq=s1
						print "star as srting", str(starseq)
						fivep= False
						print "mirname", 


					except:	
						try:	
							print "trying: ",name3, "-3p"	
							maturename=maturesdict[name3+"-3p"].id
							print "Try;",maturename
							m1=str(maturesdict[name3+"-3p"].seq)
							print "seq;",m1
							starname=starsdict[name3+"-5p*"].id
							starname
							s1=str(starsdict[name3+"-5p*"].seq)
							fivep= False
						except:
							try:
								print "trying: ",name3, "_5p"
								print  maturesdict[name3+"_5p"].id
								maturename=maturesdict[name3+"_5p"].id
								print "Try;",maturename
								m1=str(maturesdict[name3+"_5p"].seq)
								starname=starsdict[name3+"_3p*"].id
								s1=str(starsdict[name3+"_3p*"].seq)
								fivep=True

							except:
								print "trying: ",name3, "-5p"
								print  maturesdict[name3+"-5p"].id
								maturename=maturesdict[name3+"-5p"].id
								print "Try;",maturename
								m1=str(maturesdict[name3+"-5p"].seq)
								starname=starsdict[name3+"-3p*"].id
								s1=str(starsdict[name3+"-3p*"].seq)
								fivep=True



					line_us=line.replace("T", "U")

					if  "5P" in maturename or "5p" in maturename :# if 5p

						precin=line.find(m1)
						precend=line.find(s1)
				
						prec_seq_flanking_bases=line[precin-flanking_bases:precend+len(s1)+flanking_bases]
						prec_seq_flanking_bases_us=line_us[precin-flanking_bases:precend+len(s1)+flanking_bases]

						outfile.write(prec_seq_flanking_bases_us)

						outfile.close()


					else:# if 3p
						print "prec", line
						print "matseq", m1
						print "starseq", s1
						precin=line.find(s1)
						print "starfound"
						precend=line.find(m1)
						print"maturefound"


						prec_seq_flanking_bases=line[precin-flanking_bases:precend+len(m1)+flanking_bases]
						prec_seq_flanking_bases_us=line_us[precin-flanking_bases:precend+len(m1)+flanking_bases]

						outfile.write(prec_seq_flanking_bases_us)


						outfile.close()


					prec_1=prec_seq_flanking_bases.split(m1)[0]### precursor until mature
					prec_2=prec_seq_flanking_bases.split(s1)[0]### precursor until star


					if  "5P" in maturename or "5p" in maturename :# if 5'== mature
						try:
							fivep= True
							paragraph.add_run(prec_1)## from begining until mature
							run2=paragraph.add_run(m1)# mature	
							font2=run2.font
							font2.bold=True
							font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
							paragraph.add_run(prec_seq_flanking_bases.split(m1)[1].split(s1)[0])### prec from end mature until star
							run3=paragraph.add_run(s1)	
							font3=run3.font
							font3.bold=True	
							font3.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	

							paragraph.add_run(prec_seq_flanking_bases.split(s1)[1]+"\n")# from final star to end
						except:
							paragraph.add_run("\nERROR X1: strange mature/star,"+ prec_1+"ups..."+prec_seq_flanking_bases+"XXXXXX"+m1+"yyyyyyy"+str(precin)+" AAA  "+str(precend)+"\n")
					else: ## mature=3'
						try:
							fivep = False
							paragraph.add_run(prec_2)## from begining until star
							run3=paragraph.add_run(s1)	#star
							font3=run3.font
							font3.bold=True
							font3.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
							paragraph.add_run(prec_seq_flanking_bases.split(s1)[1].split(m1)[0])### prec from end star until mature
							run2=paragraph.add_run(m1)# mature	
							font2=run2.font
							font2.bold=True
							font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)	


							paragraph.add_run(prec_seq_flanking_bases.split(m1)[1]+"\n")# from final mature to end
						except:
							paragraph.add_run("ERROR 2: strange mature/star")

				#print mature
					mature_write=">"+maturename+"\n"
					mature_write_2=m1
					paragraph.add_run(mature_write).bold = True	
					run2=paragraph.add_run(mature_write_2+"\n\n")	
					font2=run2.font
					font2.bold=True
					font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)	

				#print star
					star_write=">"+starname+"*"+"\n"
					star_write_2=s1
					paragraph.add_run(star_write).bold = True	
					run1=paragraph.add_run(star_write_2+"\n\n")	
					font1=run1.font
					font1.color.rgb = RGBColor(0x1B, 0x01, 0xFF)
					font1.bold=True			
			


					
					call=str("./scripts/mfold SEQ="+name_file)
					print "\nIS calling:: ", call,"\n"
					os.system(call)
					########### Transform vienna format			
					callformatingvienna=str("perl ct2b.pl "+name_file+".ct > "+name_file+".vienna.txt")
					os.system(callformatingvienna)
					os.system("rm -rf *.fa")
					os.system("rm -rf *.det")		
					os.system("rm -rf *.plot")
					os.system("rm -rf *.ct")				
					os.system("rm -rf *.cs")
					os.system("rm -rf *.ss")
					os.system("rm -rf *.ps")
					os.system("rm -rf *.sav")
					os.system("rm -rf *.ann")
					os.system("rm -rf *-count")
					os.system("rm -rf *.pnt")			
					os.system("rm -rf *.log")			
					os.system("rm -rf *-num")
					os.system("rm -rf *.pdf")
					numline=0
	


					os.system("cat "+name_file+".out >> structures.txt")
			
					### import vienna output format
					with  open(str(name_file+".vienna.txt"), 'r') as viennafile:
						vienna= viennafile.readlines()
					viennafile.closed

					###### ADD VIENNA FORMAT
					try:	
						vienna_0=vienna[0].strip()
						vienna_1=vienna[1].split(" ")[0].strip()
					except: 
						a=1

					#### COLOR VIENNA 
					if fivep == False:# if 3' mature

						try:
							star_u=s1.replace("T", "U").rstrip()
							mature_u=m1.replace("T", "U").rstrip()
				

							prestar=vienna_0.split(star_u)[0]

							loop=vienna_0.split(star_u)[1].split(mature_u)[0]



							aftermature=vienna_0.split(mature_u)[1]

							pre_formated=paragraph.add_run(prestar)
							pre_formated.font.size=Pt(7)
				
							star_formated=paragraph.add_run(star_u)
							star_formated.font.size=Pt(7)
							font4=star_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	

							loop_formated=paragraph.add_run(loop)
							loop_formated.font.size=Pt(7)
				
							mature_formated=paragraph.add_run(mature_u)
							mature_formated.font.size=Pt(7)
							font4=mature_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)

							aftermature_formated=paragraph.add_run(aftermature)
							aftermature_formated.font.size=Pt(7)
							paragraph.add_run("\n")
							vi2=paragraph.add_run(vienna[1])
							vi2.font.size=Pt(7)
							paragraph.add_run("\n")

							####### count paired bases inside miRNA
							pren_corresponing_star=vienna_1[len(prestar):len(prestar)+len(star_u)]
							pren_corresponing_mature=vienna_1[len(prestar)+len(star_u)+len(loop):len(prestar)+len(star_u)+len(loop)+len(mature_u)]

							real_paired=0
				
							prentesis_counter=vienna_1.count("(")
							prentesis_before_star=vienna_1[0:len(prestar)].count("(")
							prentesis_after_star=vienna_1[len(prestar)+len(star_u):len(vienna_1)].count("(")
							prentesis_in_star=range(prentesis_before_star+1, prentesis_counter-prentesis_after_star+1 )## +1 inici beacause inclusive and 1 end because is exclusive


							prentesis_reverse=vienna_1.count(")")
							prentesis_after_mature=vienna_1[len(vienna_1)-len(aftermature):len(vienna_1)].count(")")
							prentesis_before_mature=vienna_1[0:len(prestar)+len(star_u)+len(loop)].count(")")
							prentesis_before_mature=vienna_1[0:len(prestar)+len(star_u)+len(loop)].count(")")
				
							prentesis_in_mature=range( prentesis_reverse-(prentesis_reverse-prentesis_after_mature)+1, (prentesis_reverse-prentesis_before_mature)+1)## +1 inici beacause inclusive and 1 end because is exclusive
				
			 					
							real_paired=set(prentesis_in_mature).intersection(prentesis_in_star)

							formated1=paragraph.add_run("Matched nts: "+ str(len(real_paired)))
							font4=formated1.font
							font4.bold=True	
							paragraph.add_run("\n")
							### Do new gff3 file				
							#for anot in annotationfile:
							#	if  precursorname.split(">")[1].split("-precursor")[0]+"-" in anot:
							#		Candidates_filter.write(anot)

			
							#for info in candidate_counts:
							#	if  precursorname.split(">")[1].split("-precursor")[0]+"-" in info:
							#		controlscounts="Mature counts in control: "+str(int(int(info.split(",")[1])+(int(info.split(",")[2]))))
							#		a=paragraph.add_run("\n"+controlscounts+"\t")
							#		a.font.size=Pt(8)
							#		starscounts="Star counts in control: "+str(int(int(info.split(",")[5])+(int(info.split(",")[6]))))
							#		a=paragraph.add_run(starscounts+"\n"+"\n")
							#		a.font.size=Pt(8)
						except:
							paragraph.add_run("ERRor 1")

					#### COLOR VIENNA 
					if fivep == True:# if 5' mature
						try:
							star_u=s1.replace("T", "U").rstrip()
							mature_u=m1.replace("T", "U").rstrip()
				
							premature=vienna_0.split(mature_u)[0]
							loop=vienna_0.split(mature_u)[1].split(star_u)[0]



							afterstar=vienna_0.split(star_u)[1]


							pre_formated=paragraph.add_run(premature)
							pre_formated.font.size=Pt(7)
				
							mature_formated=paragraph.add_run(mature_u)
							mature_formated.font.size=Pt(7)
							font4=mature_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)				
	
							loop_formated=paragraph.add_run(loop)
							loop_formated.font.size=Pt(7)
				
							star_formated=paragraph.add_run(star_u)
							star_formated.font.size=Pt(7)
							font4=star_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)
				
							afterstar_formated=paragraph.add_run(afterstar)
							afterstar_formated.font.size=Pt(7)
							paragraph.add_run("\n")
							vi2=paragraph.add_run(vienna[1])
							vi2.font.size=Pt(7)

							####### count paired bases inside miRNA
							print vienna_1
							pren_corresponing_mature=vienna_1[len(premature):len(premature)+len(mature_u)]
							pren_corresponing_star=vienna_1[len(premature)+len(mature_u)+len(loop):len(premature)+len(mature_u)+len(loop)+len(star_u)]



							paragraph.add_run("\n")

							real_paired=0
				
							prentesis_counter=vienna_1.count("(")
							prentesis_before_mature=vienna_1[0:len(premature)].count("(")
							prentesis_after_mature=vienna_1[len(premature)+len(mature_u):len(vienna_1)].count("(")
							prentesis_in_mature=range(prentesis_before_mature+1, prentesis_counter-prentesis_after_mature+1 )## +1 inici beacause inclusive and 1 end because is exclusive
							#print  prentesis_in_mature, "premature; ", prentesis_before_mature ,"aftermature ", prentesis_after_mature, vienna_1[0:len(premature)]

							prentesis_reverse=vienna_1.count(")")
							prentesis_after_star=vienna_1[len(vienna_1)-len(afterstar):len(vienna_1)].count(")")
							prentesis_before_star=vienna_1[0:len(premature)+len(mature_u)+len(loop)].count(")")
							prentesis_before_star=vienna_1[0:len(premature)+len(mature_u)+len(loop)].count(")")
				
							prentesis_in_star=range( prentesis_reverse-(prentesis_reverse-prentesis_after_star)+1, (prentesis_reverse-prentesis_before_star)+1)## +1 inici beacause inclusive and 1 end because is exclusive
				
					
							real_paired=set(prentesis_in_mature).intersection(prentesis_in_star)


							formated1=paragraph.add_run("Matched nts: "+ str(len(real_paired)))
							font4=formated1.font
							font4.bold=True	
							paragraph.add_run("\n")
						except:
							paragraph.add_run("ERROR 3: strange mature/star")

				



					os.system("rm -rf *.vienna.txt")

					mirnum=mirnum+2

					### Process mfold output and print
					try:
						with  open(str(name_file+".out"), 'r') as f:
							mfoldout= f.readlines()
						f.closed
					except:
						a=0
			
					try:
						freeenergy=mfoldout[4].split("-")[1]
						paragraph.add_run(mfoldout[4])## free energie line
						outfile_free_energies.write(str(name_file+"\t"+str(mfoldout[4])))
					except:
						paragraph.add_run("error free energie")
						freeenergy=0
				

					if 0 ==0 :## 
						try:
							paragraph.add_run(mfoldout[5])
							paragraph.add_run(mfoldout[6])
				
							topaint= True

							if  len(mfoldout) > 20:  ## if multiline folding
								if "Structure" not in mfoldout[15]:
									topaint= False
									#paragraph.add_run(mfoldout[7:25])
									paragraph.add_run(mfoldout)


							if fivep == False and topaint== True:# if 3' mature
								try:
									line1= mfoldout[7] ## first 2 lines
									line2= mfoldout[8]

									slash_i_1=line1[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line2[0:flanking_bases].count("-")
				
									slash_mid_1=line1[flanking_bases:flanking_bases+len(s1)].count("-")
									slash_mid_1=slash_mid_1+line2[flanking_bases:flanking_bases+len(s1)].count("-")

									starseq=line1[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)] # from flanking_bases + slash to flanking_bases+slash+leng of star

									before_star=line1.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
									after_star=line1.split(starseq)[1]
									paragraph.add_run(after_star)

									starseq=line2[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)]

									before_star=line2.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)
									after_star=line2.split(starseq)[1]
									paragraph.add_run(after_star)



									line3= mfoldout[9] ## second 2 lines
									line4= mfoldout[10]


									slash_i_1=line3[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line4[0:flanking_bases].count("-")
				
									slash_mid_1=line3[flanking_bases:flanking_bases+len(m1)].count("-")
									slash_mid_1=slash_mid_1+line4[flanking_bases:flanking_bases+len(m1)].count("-")

									matureseq=line3[flanking_bases+slash_i_1:flanking_bases+slash_i_1+slash_mid_1+len(m1)] # from flanking_bases + slash to flanking_bases+slash+leng of mature

									before_mature=line3.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line3.split(matureseq)[1]
									paragraph.add_run(after_mature)

									matureseq=line4[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(m1)]

									before_mature=line4.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line4.split(matureseq)[1]
									paragraph.add_run(after_mature)
								except:

									paragraph.add_run("Error XXXX")

							elif fivep == True and topaint== True:# if 5' mature
								try:

									line1= mfoldout[7] ## first 2 lines
									line2= mfoldout[8]

									slash_i_1=line1[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line2[0:flanking_bases].count("-")
				
									slash_mid_1=line1[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(m1)].count("-")
									slash_mid_1=slash_mid_1+line2[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(m1)].count("-")

									matureseq=line1[flanking_bases+slash_i_1:flanking_bases+slash_i_1+slash_mid_1+len(m1)] # from flanking_bases + slash to flanking_bases+slash+leng of mature

									before_mature=line1.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line1.split(matureseq)[1]
									paragraph.add_run(after_mature)

									matureseq=line2[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(m1)]

									before_mature=line2.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line2.split(matureseq)[1]
									paragraph.add_run(after_mature)



									line3= mfoldout[9] ## second 2 lines
									line4= mfoldout[10]

									slash_i_1=line3[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line4[0:flanking_bases].count("-")
				
									slash_mid_1=line3[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(s1)].count("-")
									slash_mid_1=slash_mid_1+line4[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(s1)].count("-")

									starseq=line3[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)] # from flanking_bases + slash to flanking_bases+slash+leng of star

									before_star=line3.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
									after_star=line3.split(starseq)[1]
									paragraph.add_run(after_star)

									starseq=line4[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)]

									before_star=line4.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
									after_star=line4.split(starseq)[1]
									paragraph.add_run(after_star)
								except:
									paragraph.add_run("ERROR YYYY")	

							#paragraph.add_run(mfoldout[flanking_bases])	



							if len(mfoldout)>22 and "Structure" not in mfoldout[15]:
								paragraph.add_run(mfoldout[12:21])
							else:
								paragraph.add_run(mfoldout[11])


						except:
							paragraph.add_run("ERROR big error")	





					document.add_picture( pictname, width=Inches(5.0) )
					paragraph.add_run("\n-----------------------------------------------------")
					document.add_page_break()


				os.system("rm -rf Homolog*")
				os.system("rm -rf *.out")

		document.save(output_dir+'miRNA_Report.docx')		
		outfile_free_energies.close()
		os.system("rm -rf precursors.gff3")
		os.system("rm -rf matures.gff3")			
		os.system("rm -rf stars.gff3")			
		os.system("rm -rf precursors")			
		os.system("rm -rf matures")			
		os.system("rm -rf stars")			
		os.system("rm -rf precursors_+flanking_bases.gff3")			
		os.system("rm -rf  precursors_plus_flanking_bases_online")			
		os.system("rm -rf  precursors_plus_flanking_bases")			
		os.system("rm -rf *.fa-local.seq")
		os.system("rm -rf *.fa.cmd")
def doreportprecsbam(comment):
		print comment
		print("""
			#####################################################################
			#####################################################################
			############						 ############
			############	 Start Folding				 ############
			############						 ############
			#####################################################################
			#####################################################################

			"""	)



		os.system("rm -rf temp/structures.txt")



		document = Document()

		document.add_heading('miRNA check', 0)

		style=document.styles["Normal"]
		font=style.font
		font.name="Mono"
		font.size=Pt(8)


		cmd="gffread -g "+fastaref+" "+gff3Matures +" -w temp/matures"
		print(cmd)
		os.system(cmd)


		cmd="gffread -g "+fastaref+" "+gff3Stars+" -w temp/stars"
		print(cmd)
		os.system(cmd)

		input_file = open("temp/matures")
		maturesdict = SeqIO.to_dict(SeqIO.parse(input_file, "fasta"))
		input_file = open("temp/stars")
		starsdict = SeqIO.to_dict(SeqIO.parse(input_file, "fasta"))



		mirnum=1






		with open(fastaref, 'r') as inF:
			for line in inF:
				if ">" in line:
					print "Line", line
					precursorname=line.rstrip()

					pictname=output_dir+"plots/"+precursorname.split(">")[1]+'.png'

					if os.path.isfile(pictname) == False:### if picture does ot exist, next candidaate
						print "doesnt exist", pictname 
						flag="next!"
						mirnum=mirnum+2
					else:## if pict exists, continue
						print "exists", pictname
						flag="continue"


						paragraph = document.add_paragraph()
						prename=line.split("_pre")[0]
						name=prename.split(">")[1].rstrip()
						name2=prename.split(">")[1].rstrip()
						name3=name.split("_pre")[0]
						print "1 - This is the name file: ", name
						name_file=name+".fa"
						outfile= open(name_file, 'w')
						outfile.write(line)

						paragraph.add_run(precursorname+"\n").bold = True
			


				if ">" not in line and flag=="continue":
					print "Name AA;", name3
					
					try:		
						maturename=maturesdict[name3+"_3p"].id
						print "Try;",maturename
						m1=str(maturesdict[name3+"_3p"].seq)
						print "seq;",m1
						starname=starsdict[name3+"_5p*"].id
						s1=str(starsdict[name3+"_5p*"].seq)
						print "star",s1
						starseq=s1
						print "star as srting", str(starseq)
						fivep= False
						print "mirname", 


					except:	
						try:	
							print "trying: ",name3, "-3p"	
							maturename=maturesdict[name3+"-3p"].id
							print "Try;",maturename
							m1=str(maturesdict[name3+"-3p"].seq)
							print "seq;",m1
							starname=starsdict[name3+"-5p*"].id
							starname
							s1=str(starsdict[name3+"-5p*"].seq)
							fivep= False
						except:
							try:
								print "trying: ",name3, "_5p"
								print  maturesdict[name3+"_5p"].id
								maturename=maturesdict[name3+"_5p"].id
								print "Try;",maturename
								m1=str(maturesdict[name3+"_5p"].seq)
								starname=starsdict[name3+"_3p*"].id
								s1=str(starsdict[name3+"_3p*"].seq)
								fivep=True

							except:
								print "trying: ",name3, "-5p"
								print  maturesdict[name3+"-5p"].id
								maturename=maturesdict[name3+"-5p"].id
								print "Try;",maturename
								m1=str(maturesdict[name3+"-5p"].seq)
								starname=starsdict[name3+"-3p*"].id
								s1=str(starsdict[name3+"-3p*"].seq)
								fivep=True



					line_us=line.replace("T", "U")

					if  "5P" in maturename or "5p" in maturename :# if 5p

						precin=line.find(m1)
						precend=line.find(s1)
				
						prec_seq_flanking_bases=line[precin-flanking_bases:precend+len(s1)+flanking_bases]
						prec_seq_flanking_bases_us=line_us[precin-flanking_bases:precend+len(s1)+flanking_bases]
						print " A Sequence on file:", prec_seq_flanking_bases_us

						outfile.write(prec_seq_flanking_bases_us)
						outfile.close()


					else:# if 3p
						precin=line.find(s1)
						precend=line.find(m1)


						prec_seq_flanking_bases=line[precin-flanking_bases:precend+len(m1)+flanking_bases]
						prec_seq_flanking_bases_us=line_us[precin-flanking_bases:precend+len(m1)+flanking_bases]

						print " B Sequence on file:", prec_seq_flanking_bases_us
						outfile.write(prec_seq_flanking_bases_us)
						outfile.close()


					prec_1=prec_seq_flanking_bases.split(m1)[0]### precursor until mature
					prec_2=prec_seq_flanking_bases.split(s1)[0]### precursor until star


					if  "5P" in maturename or "5p" in maturename :# if 5'== mature
						try:
							fivep= True
							paragraph.add_run(prec_1)## from begining until mature
							run2=paragraph.add_run(m1)# mature	
							font2=run2.font
							font2.bold=True
							font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
							paragraph.add_run(prec_seq_flanking_bases.split(m1)[1].split(s1)[0])### prec from end mature until star
							run3=paragraph.add_run(s1)	
							font3=run3.font
							font3.bold=True	
							font3.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	

							paragraph.add_run(prec_seq_flanking_bases.split(s1)[1]+"\n")# from final star to end
						except:
							paragraph.add_run("\nERROR X1: strange mature/star,"+ prec_1+"ups..."+prec_seq_flanking_bases+"XXXXXX"+m1+"yyyyyyy"+str(precin)+" AAA  "+str(precend)+"\n")
					else: ## mature=3'
						try:
							fivep = False
							paragraph.add_run(prec_2)## from begining until star
							run3=paragraph.add_run(s1)	#star
							font3=run3.font
							font3.bold=True
							font3.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
							paragraph.add_run(prec_seq_flanking_bases.split(s1)[1].split(m1)[0])### prec from end star until mature
							run2=paragraph.add_run(m1)# mature	
							font2=run2.font
							font2.bold=True
							font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)	


							paragraph.add_run(prec_seq_flanking_bases.split(m1)[1]+"\n")# from final mature to end
						except:
							paragraph.add_run("ERROR 2: strange mature/star")

				#print mature
					mature_write=">"+maturename+"\n"
					mature_write_2=m1
					paragraph.add_run(mature_write).bold = True	
					run2=paragraph.add_run(mature_write_2+"\n\n")	
					font2=run2.font
					font2.bold=True
					font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)	

				#print star
					star_write=">"+starname+"*"+"\n"
					star_write_2=s1
					paragraph.add_run(star_write).bold = True	
					run1=paragraph.add_run(star_write_2+"\n\n")	
					font1=run1.font
					font1.color.rgb = RGBColor(0x1B, 0x01, 0xFF)
					font1.bold=True			
			



					call=str("./scripts/mfold SEQ="+name_file)
					print "\nIS calling:: ", call,"\n"
					os.system(call)
					########### Transform vienna format			
					callformatingvienna=str("perl ct2b.pl "+name_file+".ct > "+name_file+".vienna.txt")
					os.system(callformatingvienna)
					os.system("rm -rf *.fa")
					os.system("rm -rf *.det")		
					os.system("rm -rf *.plot")
					os.system("rm -rf *.ct")				
					os.system("rm -rf *.cs")
					os.system("rm -rf *.ss")
					os.system("rm -rf *.ps")
					os.system("rm -rf *.sav")
					os.system("rm -rf *.ann")
					os.system("rm -rf *-count")
					os.system("rm -rf *.pnt")			
					os.system("rm -rf *.log")			
					os.system("rm -rf *-num")
					os.system("rm -rf *.pdf")
					numline=0
	


					os.system("cat "+name_file+".out >> structures.txt")
			
					### import vienna output format
					with  open(str(name_file+".vienna.txt"), 'r') as viennafile:
						vienna= viennafile.readlines()
					viennafile.closed

					###### ADD VIENNA FORMAT
					try:	
						vienna_0=vienna[0].strip()
						vienna_1=vienna[1].split(" ")[0].strip()
					except: 
						a=1

					#### COLOR VIENNA 
					if fivep == False:# if 3' mature

						try:
							star_u=s1.replace("T", "U").rstrip()
							mature_u=m1.replace("T", "U").rstrip()
				

							prestar=vienna_0.split(star_u)[0]

							loop=vienna_0.split(star_u)[1].split(mature_u)[0]



							aftermature=vienna_0.split(mature_u)[1]

							pre_formated=paragraph.add_run(prestar)
							pre_formated.font.size=Pt(7)
				
							star_formated=paragraph.add_run(star_u)
							star_formated.font.size=Pt(7)
							font4=star_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	

							loop_formated=paragraph.add_run(loop)
							loop_formated.font.size=Pt(7)
				
							mature_formated=paragraph.add_run(mature_u)
							mature_formated.font.size=Pt(7)
							font4=mature_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)

							aftermature_formated=paragraph.add_run(aftermature)
							aftermature_formated.font.size=Pt(7)
							paragraph.add_run("\n")
							vi2=paragraph.add_run(vienna[1])
							vi2.font.size=Pt(7)
							paragraph.add_run("\n")

							####### count paired bases inside miRNA
							pren_corresponing_star=vienna_1[len(prestar):len(prestar)+len(star_u)]
							pren_corresponing_mature=vienna_1[len(prestar)+len(star_u)+len(loop):len(prestar)+len(star_u)+len(loop)+len(mature_u)]

							real_paired=0
				
							prentesis_counter=vienna_1.count("(")
							prentesis_before_star=vienna_1[0:len(prestar)].count("(")
							prentesis_after_star=vienna_1[len(prestar)+len(star_u):len(vienna_1)].count("(")
							prentesis_in_star=range(prentesis_before_star+1, prentesis_counter-prentesis_after_star+1 )## +1 inici beacause inclusive and 1 end because is exclusive


							prentesis_reverse=vienna_1.count(")")
							prentesis_after_mature=vienna_1[len(vienna_1)-len(aftermature):len(vienna_1)].count(")")
							prentesis_before_mature=vienna_1[0:len(prestar)+len(star_u)+len(loop)].count(")")
							prentesis_before_mature=vienna_1[0:len(prestar)+len(star_u)+len(loop)].count(")")
				
							prentesis_in_mature=range( prentesis_reverse-(prentesis_reverse-prentesis_after_mature)+1, (prentesis_reverse-prentesis_before_mature)+1)## +1 inici beacause inclusive and 1 end because is exclusive
				
			 					
							real_paired=set(prentesis_in_mature).intersection(prentesis_in_star)

							formated1=paragraph.add_run("Matched nts: "+ str(len(real_paired)))
							font4=formated1.font
							font4.bold=True	
							paragraph.add_run("\n")


						except:
							paragraph.add_run("ERRor 1")

					#### COLOR VIENNA 
					if fivep == True:# if 5' mature
						try:
							star_u=s1.replace("T", "U").rstrip()
							mature_u=m1.replace("T", "U").rstrip()
				
							premature=vienna_0.split(mature_u)[0]
							loop=vienna_0.split(mature_u)[1].split(star_u)[0]



							afterstar=vienna_0.split(star_u)[1]


							pre_formated=paragraph.add_run(premature)
							pre_formated.font.size=Pt(7)
				
							mature_formated=paragraph.add_run(mature_u)
							mature_formated.font.size=Pt(7)
							font4=mature_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)				
	
							loop_formated=paragraph.add_run(loop)
							loop_formated.font.size=Pt(7)
				
							star_formated=paragraph.add_run(star_u)
							star_formated.font.size=Pt(7)
							font4=star_formated.font
							font4.bold=True
							font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)
				
							afterstar_formated=paragraph.add_run(afterstar)
							afterstar_formated.font.size=Pt(7)
							paragraph.add_run("\n")
							vi2=paragraph.add_run(vienna[1])
							vi2.font.size=Pt(7)

							####### count paired bases inside miRNA
							print vienna_1
							pren_corresponing_mature=vienna_1[len(premature):len(premature)+len(mature_u)]
							pren_corresponing_star=vienna_1[len(premature)+len(mature_u)+len(loop):len(premature)+len(mature_u)+len(loop)+len(star_u)]



							paragraph.add_run("\n")

							real_paired=0
				
							prentesis_counter=vienna_1.count("(")
							prentesis_before_mature=vienna_1[0:len(premature)].count("(")
							prentesis_after_mature=vienna_1[len(premature)+len(mature_u):len(vienna_1)].count("(")
							prentesis_in_mature=range(prentesis_before_mature+1, prentesis_counter-prentesis_after_mature+1 )## +1 inici beacause inclusive and 1 end because is exclusive
							#print  prentesis_in_mature, "premature; ", prentesis_before_mature ,"aftermature ", prentesis_after_mature, vienna_1[0:len(premature)]

							prentesis_reverse=vienna_1.count(")")
							prentesis_after_star=vienna_1[len(vienna_1)-len(afterstar):len(vienna_1)].count(")")
							prentesis_before_star=vienna_1[0:len(premature)+len(mature_u)+len(loop)].count(")")
							prentesis_before_star=vienna_1[0:len(premature)+len(mature_u)+len(loop)].count(")")
				
							prentesis_in_star=range( prentesis_reverse-(prentesis_reverse-prentesis_after_star)+1, (prentesis_reverse-prentesis_before_star)+1)## +1 inici beacause inclusive and 1 end because is exclusive
				
					
							real_paired=set(prentesis_in_mature).intersection(prentesis_in_star)


							formated1=paragraph.add_run("Matched nts: "+ str(len(real_paired)))
							font4=formated1.font
							font4.bold=True	
							paragraph.add_run("\n")
						except:
							paragraph.add_run("ERROR 3: strange mature/star")

				



					os.system("rm -rf *.vienna.txt")

					mirnum=mirnum+2

					### Process mfold output and print
					try:
						with  open(str(name_file+".out"), 'r') as f:
							mfoldout= f.readlines()
						f.closed
					except:
						a=0
			
					try:
						freeenergy=mfoldout[4].split("-")[1]
						paragraph.add_run(mfoldout[4])## free energie line
						outfile_free_energies.write(str(name_file+"\t"+str(mfoldout[4])))
					except:
						paragraph.add_run("error free energie")
						freeenergy=0
				

					if 0 ==0 :## 
						try:
							paragraph.add_run(mfoldout[5])
							paragraph.add_run(mfoldout[6])
				
							topaint= True

							if  len(mfoldout) > 20:  ## if multiline folding
								if "Structure" not in mfoldout[15]:
									topaint= False
									#paragraph.add_run(mfoldout[7:25])
									paragraph.add_run(mfoldout)


							if fivep == False and topaint== True:# if 3' mature
								try:
									line1= mfoldout[7] ## first 2 lines
									line2= mfoldout[8]

									slash_i_1=line1[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line2[0:flanking_bases].count("-")
				
									slash_mid_1=line1[flanking_bases:flanking_bases+len(s1)].count("-")
									slash_mid_1=slash_mid_1+line2[flanking_bases:flanking_bases+len(s1)].count("-")

									starseq=line1[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)] # from flanking_bases + slash to flanking_bases+slash+leng of star

									before_star=line1.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
									after_star=line1.split(starseq)[1]
									paragraph.add_run(after_star)

									starseq=line2[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)]

									before_star=line2.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)
									after_star=line2.split(starseq)[1]
									paragraph.add_run(after_star)



									line3= mfoldout[9] ## second 2 lines
									line4= mfoldout[10]


									slash_i_1=line3[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line4[0:flanking_bases].count("-")
				
									slash_mid_1=line3[flanking_bases:flanking_bases+len(m1)].count("-")
									slash_mid_1=slash_mid_1+line4[flanking_bases:flanking_bases+len(m1)].count("-")

									matureseq=line3[flanking_bases+slash_i_1:flanking_bases+slash_i_1+slash_mid_1+len(m1)] # from flanking_bases + slash to flanking_bases+slash+leng of mature

									before_mature=line3.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line3.split(matureseq)[1]
									paragraph.add_run(after_mature)

									matureseq=line4[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(m1)]

									before_mature=line4.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line4.split(matureseq)[1]
									paragraph.add_run(after_mature)
								except:

									paragraph.add_run("Error XXXX")

							elif fivep == True and topaint== True:# if 5' mature
								try:

									line1= mfoldout[7] ## first 2 lines
									line2= mfoldout[8]

									slash_i_1=line1[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line2[0:flanking_bases].count("-")
				
									slash_mid_1=line1[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(m1)].count("-")
									slash_mid_1=slash_mid_1+line2[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(m1)].count("-")

									matureseq=line1[flanking_bases+slash_i_1:flanking_bases+slash_i_1+slash_mid_1+len(m1)] # from flanking_bases + slash to flanking_bases+slash+leng of mature

									before_mature=line1.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line1.split(matureseq)[1]
									paragraph.add_run(after_mature)

									matureseq=line2[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(m1)]

									before_mature=line2.split(matureseq)[0]
									paragraph.add_run(before_mature)
									run4=paragraph.add_run(matureseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
									after_mature=line2.split(matureseq)[1]
									paragraph.add_run(after_mature)



									line3= mfoldout[9] ## second 2 lines
									line4= mfoldout[10]

									slash_i_1=line3[0:flanking_bases].count("-")
									slash_i_1=slash_i_1+line4[0:flanking_bases].count("-")
				
									slash_mid_1=line3[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(s1)].count("-")
									slash_mid_1=slash_mid_1+line4[flanking_bases+slash_i_1:flanking_bases+slash_i_1+len(s1)].count("-")

									starseq=line3[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)] # from flanking_bases + slash to flanking_bases+slash+leng of star

									before_star=line3.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
									after_star=line3.split(starseq)[1]
									paragraph.add_run(after_star)

									starseq=line4[flanking_bases+slash_i_1 : flanking_bases+slash_i_1+slash_mid_1+len(s1)]

									before_star=line4.split(starseq)[0]
									paragraph.add_run(before_star)
									run4=paragraph.add_run(starseq)
									font4=run4.font
									font4.bold=True	
									font4.color.rgb = RGBColor(0x1B, 0x01, 0xFF)	
									after_star=line4.split(starseq)[1]
									paragraph.add_run(after_star)
								except:
									paragraph.add_run("ERROR YYYY")	

							#paragraph.add_run(mfoldout[flanking_bases])	



							if len(mfoldout)>22 and "Structure" not in mfoldout[15]:
								paragraph.add_run(mfoldout[12:21])
							else:
								paragraph.add_run(mfoldout[11])


						except:
							paragraph.add_run("ERROR big error")	





					document.add_picture( pictname, width=Inches(5.0) )
					paragraph.add_run("\n-----------------------------------------------------")
					document.add_page_break()


				os.system("rm -rf Homolog*")
				os.system("rm -rf *.out")

		document.save('miRNA_Report.docx')		

		os.system("rm -rf precursors.gff3")
		os.system("rm -rf matures.gff3")			
		os.system("rm -rf stars.gff3")			
		os.system("rm -rf precursors")			
		os.system("rm -rf matures")			
		os.system("rm -rf stars")			
		os.system("rm -rf precursors_+flanking_bases.gff3")			
		os.system("rm -rf  precursors_plus_flanking_bases_online")			
		os.system("rm -rf  precursors_plus_flanking_bases")			
		os.system("rm -rf *.fa-local.seq")
		os.system("rm -rf *.fa.cmd")
		os.system("rm -rf *.fa.out")
###########################################################################################################################################

version = '1.0'
verbose = False
output_dir = 'Figures'
OnlyPngPlot=False
flanking_bases=5

options, remainder = getopt.getopt(sys.argv[1:], 'o:vh', [   							'output=', 
														 'verbose',
														 'help',
														 'version=',
														 'BamFilesDir=',
														 'Genome=',
														 'gff3Precs=',
														 'gff3Matures=',
														 'gff3Stars=',
														 'PrecsFasta=',
														 'Flanking_bases=',
														 'OnlyPngPlot=',
														 ])

print 'ARGV	  :', sys.argv[1:]
print 'OPTIONS   :', options
print 'VERSION   :', version
print 'VERBOSE   :', verbose
print 'OUTPUT	:', output_dir
print 'REMAINING :', remainder
print 'Only plot :', OnlyPngPlot
#print 'BamFilesDir :', BamFilesDir


for opt, arg in options:
	if opt in ('-h', '--help'):
		print """ 
That's the help page (-h / --help)

You have differnt modes to run mirPLOT.py

 1- With Bam files of shortRNA reads mapped against miRNA precursors+some flanking nucleotides:

   1.1 - Providing precursors fasta file and bam files directory

	 $ python mirPLOT_v04.py --BamFilesDir Example_Files/bams/ --PrecsFasta Example_Files/conserved+30.fa --output Figures_2  --Flanking_bases 11

   1.2 - Providing precursors fasta file, bam files directory and gff3 file with mature miRNA annotations

	 $ python mirPLOT_v04.py --BamFilesDir Example_Files/bams/ --PrecsFasta Example_Files/conserved+30.fa --output Figures_2  --gff3Matures Example_Files/Coordinates_matures_30.gff3 --Flanking_bases 11

   1.3 - Providing precursors fasta file, bam files directory and gff3 file with mature and star miRNA annotations
	 
	 $ python mirPLOT_current.py --BamFilesDir Example_Files/bams/ --PrecsFasta Example_Files/conserved+30.fa --output Example_Files/  --gff3Matures Example_Files/Coordinates_matures_30.gff3  --gff3Stars Example_Files/Coordinates_stars_30.gff3 --Flanking_bases 11



--BamFilesDir
--PrecsFasta
--output / -o
--help / -h
--gff3Matures
--gff3Stars
--OnlyPngPlot
--Flanking_bases

contact: guille.ylla@ibe.upf-csic.es

\n"""
		sys.exit()


	elif opt in ('-v', '--verbose'):
		verbose = True

	elif opt in ('-o', '--output'):
		output_dir = arg
		if not os.path.exists(output_dir):
			os.makedirs(output_dir)
		if not os.path.exists(output_dir+"plots"):
			os.makedirs(output_dir+"plots")
		outfile_free_energies= open(output_dir+"free_energies.csv", 'w')

	elif opt == '--version':
		version = arg

	elif opt == '--OnlyPngPlot':
		OnlyPngPlot = True

	elif opt == '--BamFilesDir':
		BamFilesDir = arg
		#print os.listdir(BamFilesDir)
		for file in os.listdir(BamFilesDir):
			if file.endswith(".bam"):
				print(file)

	elif opt == '--Genome':
		Genome = arg
		if (os.path.isfile(Genome) ==False):
			print "Error ",Genome, "Not found"
			break
		else:
			print "Genome file=", Genome


	elif opt == '--gff3Precs':
		gff3Precs = arg
		if (os.path.isfile(gff3Precs) ==False):
			print "Error ",gff3Precs, "Not found"
			break
		else:
			print "gff3Precs file=", gff3Precs

	elif opt == '--gff3Matures':
		gff3Matures = arg
		if (os.path.isfile(gff3Matures) ==False):
			print "Error ",gff3Matures, "Not found"
			break
		else:
			print "gff3Matures file=", gff3Matures

	elif opt == '--gff3Stars':
		gff3Stars = arg
		if (os.path.isfile(gff3Stars) ==False):
			print "Error ",gff3Stars, "Not found"
			break
		else:
			print "gff3Stars file=", gff3Stars

	elif opt == '--Flanking_bases':
		flanking_bases=int(arg)


	elif opt == '--PrecsFasta':
		PrecsFasta = arg
		if (os.path.isfile(PrecsFasta) ==False):
			print "Error ",PrecsFasta, "Not found"
			break
	else:
		print "Missing args"

######################################################################## NO gff3
if  ( "gff3Matures" not in globals() and  "gff3Stars" not in globals() ): ## No GFF3
	os.system("Rscript scripts/Visualization_precsbam.R "+BamFilesDir+" "+ PrecsFasta+" "+output_dir+"plots")

	print("""
	#####################################################################
	#####################################################################
	############						 ############
	############	  1 Coverage Plots done!				   ############
	############						 ############
	#####################################################################
	#####################################################################

	""")
	if ( OnlyPngPlot == True):
		sys.exit()
	
	else:
		print("""
		#####################################################################
		#####################################################################
		############						 ############
		############	 start Folding without arms info	 ############
		############						 ############
		#####################################################################
		#####################################################################

		"""	)



		os.system("rm -rf temp/structures.txt")

		str(sys.argv)
		

		document = Document()

		document.add_heading('miRNA check', 0)

		style=document.styles["Normal"]
		font=style.font
		#font.name="Mono"#"Courier New"#
		font.name="Mono"
		font.size=Pt(8)


		mir=1


		with open(PrecsFasta, 'r') as inF:
			for line in inF:
				if ">" in line:
					print "Print line:", line
					precursorname=line.rstrip()

					pictname=output_dir+"plots/"+precursorname.split(">")[1]+'.png'

					if os.path.isfile(pictname) == False:### if picture does ot exist, next candidaate
						print "doesnt exist", pictname 
						flag="next!"

					else:## if pict exists, continue
						print "exists", pictname
						flag="continue"

						paragraph = document.add_paragraph()
						prename=line.split("_prec")[0]
						name=prename.split(">")[1].rstrip()
						name2=prename.split(">")[1].rstrip()
						name3=name.split("_pre")[0]
						name_file=name+".fa"
						print "2 FILE NAME:: ", name_file, line
						outfile= open(name_file, 'w+')
						outfile.write(line)

						paragraph.add_run(precursorname+"\n").bold = True

				
				if ">" not in line and flag=="continue":
					line_us=line.replace("T", "U")
					outfile.write(line_us)
					outfile.close()
	

					os.system("rm -rf *.vienna.txt")
					call=str("mfold SEQ="+name_file)
					print "\nIS calling:: ", call,"\n"
					os.system(call)
					########### Transform vienna format			
					callformatingvienna=str("perl ct2b.pl "+name_file+".ct > "+name_file+".vienna.txt")
					os.system(callformatingvienna)
					os.system("rm -rf *.fa")
					os.system("rm -rf *.det")		
					os.system("rm -rf *.plot")
					os.system("rm -rf *.ct")				
					os.system("rm -rf *.cs")
					os.system("rm -rf *.ss")
					os.system("rm -rf *.ps")
					os.system("rm -rf *.sav")
					os.system("rm -rf *.ann")
					os.system("rm -rf *-count")
					os.system("rm -rf *.pnt")			
					os.system("rm -rf *.log")			
					os.system("rm -rf *-num")
					os.system("rm -rf *.pdf")

					numline=0
	


					os.system("cat "+name_file+".out >> structures.txt")
			
					### import vienna output format
					with  open(str(name_file+".vienna.txt"), 'r') as viennafile:
						vienna= viennafile.readlines()
					viennafile.closed


					###### ADD VIENNA FORMAT	
					vienna_0=vienna[0].strip()
					vienna_1=vienna[1].split(" ")[0].strip()


		
					pre_formated=paragraph.add_run(vienna_0)
					pre_formated.font.size=Pt(6)
					paragraph.add_run("\n")
					pre_formated=paragraph.add_run(vienna_1)
					pre_formated.font.size=Pt(6)
					paragraph.add_run("\n")
					### Process mfold output and print
					with  open(str(name_file+".out"), 'r') as f:
							mfoldout= f.readlines()
					f.closed

	### print counts
		
					paragraph.add_run("\n")## free energie line
					paragraph.add_run(mfoldout[4])## free energie line
					outfile_free_energies.write(str(name_file+"\t"+str(mfoldout[4])))
					paragraph.add_run(mfoldout[5])
					paragraph.add_run(mfoldout[6])
					paragraph.add_run(mfoldout[7])
					paragraph.add_run(mfoldout[8])
					paragraph.add_run(mfoldout[9])
					paragraph.add_run(mfoldout[10])
					paragraph.add_run(mfoldout[11])

					try:					
						if "Structure" not in mfoldout[15]:
							paragraph.add_run(mfoldout[12])
							paragraph.add_run(mfoldout[13])
							paragraph.add_run(mfoldout[14])
							paragraph.add_run(mfoldout[15])
							paragraph.add_run(mfoldout[16])
							paragraph.add_run(mfoldout[17])
							paragraph.add_run(mfoldout[18])
							paragraph.add_run(mfoldout[19])
					except:
						a=1
						
			
					if(flag=="continue"):
						document.add_picture( pictname, width=Inches(5.0) )

					paragraph.add_run("\n-----------------------------------------------------")
					document.add_page_break()



		document.save(output_dir+'miRNA_structure.docx')	
		os.system("rm -rf precursors.gff3")
		os.system("rm -rf matures.gff3")			
		os.system("rm -rf stars.gff3")			
		os.system("rm -rf precursors")			
		os.system("rm -rf matures")			
		os.system("rm -rf stars")			
		os.system("rm -rf precursors_+flanking_bases.gff3")			
		os.system("rm -rf  precursors_plus_flanking_bases_online")			
		os.system("rm -rf  precursors_plus_flanking_bases")			
		os.system("rm -rf *.fa-local.seq")
		os.system("rm -rf *.fa.cmd")

######################################################################### end no gff3



if  ( "gff3Matures"  in globals() and  "gff3Stars" not in globals() ): ## mature GFF3
	os.system("Rscript scripts/Visualization_precsbam.R "+BamFilesDir+" "+ PrecsFasta+" "+output_dir+"plots"+" "+gff3Matures)

	print("""
	#####################################################################
	#####################################################################
	############						 ############
	############	2  Coverage Plots done!				   ############
	############						 ############
	#####################################################################
	#####################################################################

	""")
	if ( OnlyPngPlot == True):
		sys.exit()
	
	else:
		print("""
		#####################################################################
		#####################################################################
		############						 ############
		############	 start Folding				 ############
		############						 ############
		#####################################################################
		#####################################################################

		"""	)



		os.system("rm -rf temp/structures.txt")

		str(sys.argv)
		

		document = Document()

		document.add_heading('miRNA check', 0)

		style=document.styles["Normal"]
		font=style.font
		font.name="Mono"
		font.size=Pt(8)


		cmd="gffread -g "+PrecsFasta+" "+gff3Matures +" -w temp/matures"
		print(cmd)
		os.system(cmd)


		input_file = open("temp/matures")
		maturesdict = SeqIO.to_dict(SeqIO.parse(input_file, "fasta"))


		mir=1



		with open(PrecsFasta, 'r') as inF:
			for line in inF:
				if ">" in line:
					print "Print line:", line
					precursorname=line.rstrip()

					pictname=output_dir+"plots/"+precursorname.split(">")[1]+'.png'

					if os.path.isfile(pictname) == False:### if picture does ot exist, next candidaate
						print "doesnt exist", pictname 
						flag="next!"

					else:## if pict exists, continue
						print "exists", pictname
						flag="continue"

						paragraph = document.add_paragraph()
						prename=line.split("_prec")[0]
						name=prename.split(">")[1].rstrip()
						name2=prename.split(">")[1].rstrip()
						name3=name.split("_pre")[0]
						name_file=name+".fa"
						print "2 FILE NAME:: ", name_file, line
						outfile= open(name_file, 'w+')
						outfile.write(line)

						paragraph.add_run(precursorname+"\n").bold = True
			

				if ">" not in line and flag=="continue":
					print "NameXXXX;", name3

					try:		
						print "trying: ",name3, "_3p"
						print maturesdict[name3+"_3p"].id
						maturename=maturesdict[name3+"_3p"].id
						m1=str(maturesdict[name3+"_3p"].seq)
						print m1
						mir=maturesdict[name3+"_3p"]
						fivep= False
					except:	
						try:							
							print "trying: ",name3, "-3p"
							print  maturesdict[name3+"-3p"].id	
							maturename=maturesdict[name3+"-3p"].id
							m1=str(maturesdict[name3+"-3p"].seq)
							print m1
							mir=maturesdict[name3+"-3p"]
							fivep= False
						except:
							try:
								print "trying: ",name3, "_5p"
								print  maturesdict[name3+"_5p"].id
								maturename=maturesdict[name3+"_5p"].id
								m1=str(maturesdict[name3+"_5p"].seq)
								print m1
								mir=maturesdict[name3+"_5p"].lower()
								fivep=True

							except:
								print "trying: ",name3, "-5p"
								print  maturesdict[name3+"-5p"].id
								maturename=maturesdict[name3+"-5p"].id
								m1=str(maturesdict[name3+"-5p"].seq)
								print m1
								mir=maturesdict[name3+"-5p"]
								fivep=True

					line_us=line.replace("T", "U")

					outfile.write(line_us)
					outfile.close()

					#m1=matures[mir].rstrip()#mature

					print line
					prec_1=line.split(m1)[0]### precursor until mature
					pre_mature_len=len(line.split(m1)[0])
					post_mature_len=len(line.split(m1)[1])

					
					paragraph.add_run(prec_1)## from begining until mature
					run2=paragraph.add_run(m1)# mature	
					font2=run2.font
					font2.bold=True
					font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
					try:
						paragraph.add_run(line.split(m1)[1])### prec from end mature until star
					except:
						paragraph.add_run("ERROR 1")
			
					paragraph.add_run("\n ")

					#mature_write=matures[mir-1].split("\n")[0]+"\n"
					#mature_write_2=matures[mir]
					mature_write=">"+maturename+"\n"
					mature_write_2=m1
					paragraph.add_run(mature_write).bold = True	
					run2=paragraph.add_run(mature_write_2+"\n")	
					font2=run2.font
					font2.bold=True
					font2.color.rgb = RGBColor(0xF2, 0x09, 0x09)
					

					os.system("rm -rf *.vienna.txt")
					call=str("./scripts/mfold SEQ="+name_file)
					print "\nIS calling:: ", call,"\n"
					os.system(call)
					########### Transform vienna format			
					callformatingvienna=str("perl ct2b.pl "+name_file+".ct > "+name_file+".vienna.txt")
					os.system(callformatingvienna)
					os.system("rm -rf *.fa")
					os.system("rm -rf *.det")		
					os.system("rm -rf *.plot")
					os.system("rm -rf *.ct")				
					os.system("rm -rf *.cs")
					os.system("rm -rf *.ss")
					os.system("rm -rf *.ps")
					os.system("rm -rf *.sav")
					os.system("rm -rf *.ann")
					os.system("rm -rf *-count")
					os.system("rm -rf *.pnt")			
					os.system("rm -rf *.log")			
					os.system("rm -rf *-num")
					os.system("rm -rf *.pdf")

					numline=0
	


					os.system("cat "+name_file+".out >> structures.txt")
			
					### import vienna output format
					with  open(str(name_file+".vienna.txt"), 'r') as viennafile:
						vienna= viennafile.readlines()
					viennafile.closed


					###### ADD VIENNA FORMAT	
					vienna_0=vienna[0].strip()
					vienna_1=vienna[1].split(" ")[0].strip()

			
		
					#### COLOR VIENNA 
					mature_u=m1.replace("T", "U").rstrip()
			
					premature=vienna[0].split(mature_u)[0]
					try:
						loop=vienna[0].split(mature_u)[1]
					except:
						paragraph.add_run("ERRor 2")

					#paragraph.add_run("\n >miRNA precursor \n").bold = True
					paragraph.add_run("\n\n").bold = True
					pre_formated=paragraph.add_run(premature)
					pre_formated.font.size=Pt(6)
			
					mature_formated=paragraph.add_run(mature_u)
					mature_formated.font.size=Pt(6)
					font4=mature_formated.font
					font4.bold=True
					font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)				

					loop_formated=paragraph.add_run(loop)
					loop_formated.font.size=Pt(6)
					

					prentesis=paragraph.add_run(vienna_1)
					prentesis.font.size=Pt(6)
					paragraph.add_run("\n")

					### Process mfold output and print
					with  open(str(name_file+".out"), 'r') as f:
						mfoldout= f.readlines()
					f.closed

		### print counts
			


					if fivep == True:# if 5' mature
						try:
							paragraph.add_run("\n")## free energie line
							paragraph.add_run(mfoldout[4])## free energie line
							outfile_free_energies.write(str(name_file+"\t"+str(mfoldout[4])))
							paragraph.add_run(mfoldout[5])
							paragraph.add_run(mfoldout[6])

							line1= mfoldout[7] ## first 2 lines
							line2= mfoldout[8]


							points=line1[0:pre_mature_len].count(".")
							slash_i_1=line1[0:pre_mature_len].count("-")
							slash_i_1=slash_i_1+line2[0:pre_mature_len].count("-")
							slash_i_1=slash_i_1+points			

							slash_mid_1=line1[pre_mature_len:pre_mature_len+len(m1)].count("-")
							slash_mid_1=slash_mid_1+line2[pre_mature_len:pre_mature_len+len(m1)].count("-")

							matureseq=line1[pre_mature_len+slash_i_1:pre_mature_len+slash_i_1+slash_mid_1+len(m1)] # from 11 + slash to 11+slash+leng of mature

	
							before_mature=line1.split(matureseq)[0]	
							paragraph.add_run(before_mature)
							run4=paragraph.add_run(matureseq)
							font4=run4.font
							font4.bold=True	
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
							after_mature=line1.split(matureseq)[1]
							paragraph.add_run(after_mature)		





							matureseq=line2[pre_mature_len+slash_i_1 : pre_mature_len+slash_i_1+slash_mid_1+len(m1)]

							before_mature=line2.split(matureseq)[0]
							paragraph.add_run(before_mature)
							run4=paragraph.add_run(matureseq)
							font4=run4.font
							font4.bold=True	
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
							after_mature=line2.split(matureseq)[1]
							paragraph.add_run(after_mature)



							line3= mfoldout[9] ## second 2 lines
							line4= mfoldout[10]
							
							run4=paragraph.add_run(line3)
							run4=paragraph.add_run(line4)
							paragraph.add_run(mfoldout[11])
							try:					

								if "Structure" not in mfoldout[15]:
									paragraph.add_run(mfoldout[12])
									paragraph.add_run(mfoldout[13])
									paragraph.add_run(mfoldout[14])
									paragraph.add_run(mfoldout[15])
									paragraph.add_run(mfoldout[16])
									paragraph.add_run(mfoldout[17])
									paragraph.add_run(mfoldout[18])
									paragraph.add_run(mfoldout[19])
							except:
								a=1

						except:
							paragraph.add_run("ERROR 3: Strange structure")
							paragraph.add_run("\n")## free energie line
							paragraph.add_run(mfoldout[4])## free energie line
							outfile_free_energies.write(str(name_file+"\t"+str(mfoldout[4])))
							paragraph.add_run(mfoldout[5])
							paragraph.add_run(mfoldout[6])
							paragraph.add_run(mfoldout[7])
							paragraph.add_run(mfoldout[8])
							paragraph.add_run(mfoldout[9])
							paragraph.add_run(mfoldout[10])
							try:					
								paragraph.add_run(mfoldout[11])
								paragraph.add_run(mfoldout[12])
								paragraph.add_run(mfoldout[13])
								paragraph.add_run(mfoldout[14])
								paragraph.add_run(mfoldout[15])
								paragraph.add_run(mfoldout[16])
								paragraph.add_run(mfoldout[17])
								paragraph.add_run(mfoldout[18])
								paragraph.add_run(mfoldout[19])
							except:
								a=1

					if fivep == False:# if 3' mature


						post_mature_len=post_mature_len-1
						#paragraph.add_run("\n"+"post_mature_len ="+str(post_mature_len )+"\n")



						try:
					
							line1= mfoldout[7] ## first 2 lines
							line2= mfoldout[8]


							line3= mfoldout[9] ## second 2 lines
							line4= mfoldout[10]

							points=line3[0:post_mature_len].count(".")
							slash_i_1=line3[0:post_mature_len].count("-")
							slash_i_1=slash_i_1+line4[0:post_mature_len].count("-")
							slash_i_1=slash_i_1+points			

							slash_mid_1=line3[post_mature_len:post_mature_len+len(m1)].count("-")
							slash_mid_1=slash_mid_1+line4[post_mature_len:post_mature_len+len(m1)].count("-")

							matureseq=line3[post_mature_len+slash_i_1: post_mature_len +slash_i_1+slash_mid_1+len(m1)] # from 11 + slash to 11+slash+leng of mature

							before_mature=line3.split(matureseq)[0]	

							paragraph.add_run("\n")
							paragraph.add_run(mfoldout[4])## free energie line
							outfile_free_energies.write(str(name_file+"\t"+str(mfoldout[4])))
							paragraph.add_run(mfoldout[5])
							paragraph.add_run(mfoldout[6])
							paragraph.add_run(line1)
							paragraph.add_run(line2)
				

							paragraph.add_run(before_mature)
							run4=paragraph.add_run(matureseq)
							font4=run4.font
							font4.bold=True	
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
							after_mature=line3.split(matureseq)[1]
							paragraph.add_run(after_mature)		




				
							matureseq=line4[post_mature_len+slash_i_1 : post_mature_len+slash_i_1+slash_mid_1+len(m1)]


							before_mature=line4.split(matureseq)[0]
							paragraph.add_run(before_mature)
							run4=paragraph.add_run(matureseq)
							font4=run4.font
							font4.bold=True	
							font4.color.rgb = RGBColor(0xF2, 0x09, 0x09)	
							after_mature=line4.split(matureseq)[1]
							paragraph.add_run(after_mature)	
							paragraph.add_run(mfoldout[11])	
						except:
							paragraph.add_run("ERROR 4: Strange structure")
							paragraph.add_run("\n")## free energie line
							paragraph.add_run(mfoldout[4])## free energie line
							outfile_free_energies.write(str(name_file+"\t"+str(mfoldout[4])))
							paragraph.add_run(mfoldout[5])
							paragraph.add_run(mfoldout[6])
							paragraph.add_run(mfoldout[7])
							paragraph.add_run(mfoldout[8])
							paragraph.add_run(mfoldout[9])
							paragraph.add_run(mfoldout[10])
							paragraph.add_run(mfoldout[11])
							paragraph.add_run(mfoldout[12])
							paragraph.add_run(mfoldout[13])
							#paragraph.add_run(mfoldout[14])
							#paragraph.add_run(mfoldout[15])
							#paragraph.add_run(mfoldout[16])
							#paragraph.add_run(mfoldout[17])
							#paragraph.add_run(mfoldout[18])
							#paragraph.add_run(mfoldout[19])

			


				

					if(flag=="continue"):
						document.add_picture( pictname, width=Inches(5.0) )

					paragraph.add_run("\n-----------------------------------------------------")
					document.add_page_break()



		document.save(output_dir+'miRNA_structure.docx')	
		os.system("rm -rf precursors.gff3")
		os.system("rm -rf matures.gff3")			
		os.system("rm -rf stars.gff3")			
		os.system("rm -rf precursors")			
		os.system("rm -rf matures")			
		os.system("rm -rf stars")			
		os.system("rm -rf precursors_+flanking_bases.gff3")			
		os.system("rm -rf  precursors_plus_flanking_bases_online")			
		os.system("rm -rf  precursors_plus_flanking_bases")			
		os.system("rm -rf *.fa-local.seq")
		os.system("rm -rf *.fa.cmd")
	


if  ( "gff3Matures" in globals() and  "gff3Stars" in globals() ): # Mature and star gff3

	if "Genome" not in globals():

		os.system("Rscript scripts/Visualization_precsbam.R "+BamFilesDir+" "+PrecsFasta+" "+output_dir+"plots"+" "+gff3Matures+" "+gff3Stars)
	
		print("""
		#####################################################################
		#####################################################################
		############						 ############
		############	  Coverage Plots done!			 ############
		############						 ############
		#####################################################################
		#####################################################################

		"""	)
		fastaref=PrecsFasta

	
		if ( OnlyPngPlot == True):  ## if FALSE, do NOT FOLD
			sys.exit()
		
		else:
			doreportprecsbam("Running") #### call fucntion to do the report
	



	else:
		print "\nGENOME!!!\n"
		os.system("Rscript scripts/Visualization_genome.R "+BamFilesDir+" "+Genome+" "+output_dir+"plots"+" "+gff3Matures+" "+gff3Stars+" "+gff3Precs+" "+str(flanking_bases))
	
		print("""
		#####################################################################
		#####################################################################
		############						 ############
		############	  Coverage Plots from genome!		 ############
		############						 ############
		#####################################################################
		#####################################################################

		"""	)

		os.system("python scripts/add_extra_nts_fromgenome.py "+gff3Precs +" "+str(flanking_bases ) +" "+"> temp/precs_11.gff3")

		cmd="gffread -g "+Genome+ " temp/precs_11.gff3 -w temp/precs"
		print(cmd)
		os.system(cmd)

		os.system("perl scripts/fasta2one.pl  temp/precs > temp/precs_oneline ")


		fastaref="temp/precs_oneline"
	
		if ( OnlyPngPlot == True):  ## if FALSE, do NOT FOLD
			sys.exit()
		
		else:
			doreportgenomebam("Doing report") #### call fucntion to do the report


os.system("rm -rf *.fa.out")



